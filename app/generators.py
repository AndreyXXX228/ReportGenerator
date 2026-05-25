import os
from datetime import datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from app.models import TestRunData

REPORTS_DIR = "generated_reports"


def apply_times_new_roman(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

def apply_times_new_roman_to_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)

def _add_metrics_section(doc, data):
    doc.add_paragraph(f'Название: Отчёт по результатам автоматизированного тестирования')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Версия ПО: {data.version}')

    doc.add_paragraph('\nОсновные метрики:')
    doc.add_paragraph(f'- Всего тестов: {data.total_tests}')
    doc.add_paragraph(f'- Успешных: {data.successful_tests}')
    doc.add_paragraph(f'- С ошибкой: {data.failed_tests} ({data.error_percentage:.1f}%)')
    doc.add_paragraph(f'- Пропущенных: {data.skipped_tests} ({data.skipped_percentage:.1f}%)')

    doc.add_paragraph('\nРезультаты')

    p = doc.add_paragraph()
    if data.has_blocking_defect and data.blocking_defect:
        p.add_run('- Обнаружен блокирующий дефект:')
        bd = data.blocking_defect
        doc.add_paragraph(f'  {bd.test_case_id}, Шаг {bd.step_number}. {bd.description}')
    else:
        p.add_run('- Блокирующие дефекты отсутствуют')

    return doc

def _add_failed_tests_section(doc, data):
    p = doc.add_paragraph()
    if data.failed_test_cases:
        p.add_run(f'- {len(data.failed_test_cases)} тест-кейсов завершены с ошибкой:')
        for failed in data.failed_test_cases[:3]:
            doc.add_paragraph(f'  {failed.test_case_id}, Шаг {failed.step_number}. {failed.description}')
    else:
        p.add_run('- Все тест-кейсы завершены успешно')
    return doc

def _add_skipped_tests_section(doc, data):
    if data.skipped_test_cases:
        p = doc.add_paragraph()
        if data.has_blocking_defect:
            p.add_run(
                f'- {len(data.skipped_test_cases)} тест-кейсов пропущено, т.к. при проведении тест-кейса {data.blocking_defect.test_case_id if data.blocking_defect else "?"} обнаружена блокирующая ошибка.')
        else:
            p.add_run(
                f'- {len(data.skipped_test_cases)} тест-кейсов пропущено, т.к. тест-кейсы из предусловия завершены с ошибкой.')
        for skipped in data.skipped_test_cases[:3]:
            doc.add_paragraph(f'  {skipped.test_case_id}, Шаг {skipped.step_number}. {skipped.description}')
    return doc

def _add_decision_solution(doc, data):
    doc.add_paragraph('\nРешение:')
    p = doc.add_paragraph()
    p.add_run(f'Представленный на автоматизированное тестирование Проект «{data.project_name}»\n')
    p.add_run(f'прошел проверку в соответствии с прогоном {data.run_name} версии\n')
    p.add_run(f'{data.run_version} (приложение) и можно считать ')

    if data.failed_tests == 0 and not data.has_blocking_defect:
        p.add_run('работоспособным.')
    else:
        p.add_run('не работоспособным.')

    doc.add_paragraph()
    return doc

def _add_author_of_tests_section(doc, data):
    doc.add_paragraph()
    doc.add_paragraph('Ответственный за тестирование')
    doc.add_paragraph(data.responsible_person if data.responsible_person else '____________________')
    doc.add_paragraph(datetime.now().strftime('%d.%m.%Y'))
    return doc

def generate_summary_report(data: TestRunData) -> str:
    os.makedirs(REPORTS_DIR, exist_ok=True)

    doc = Document()
    apply_times_new_roman_to_doc(doc)

    title = doc.add_heading('Сводный отчёт', 0)
    apply_times_new_roman(title)

    doc = _add_metrics_section(doc, data)

    p = doc.add_paragraph()
    if data.failed_test_cases:
        p.add_run(f'- {len(data.failed_test_cases)} тест-кейсов завершены с ошибкой:')
        for failed in data.failed_test_cases:
            doc.add_paragraph(f'  {failed.test_case_id}, Шаг {failed.step_number}. {failed.description}')
    else:
        p.add_run('- Все тест-кейсы завершены успешно')

    if data.skipped_test_cases:
        p = doc.add_paragraph()
        if data.has_blocking_defect:
            p.add_run(
                f'- {len(data.skipped_test_cases)} тест-кейсов пропущено, т.к. при проведении тест-кейса {data.blocking_defect.test_case_id if data.blocking_defect else "?"} обнаружена блокирующая ошибка.')
        else:
            p.add_run(
                f'- {len(data.skipped_test_cases)} тест-кейсов пропущено, т.к. тест-кейсы из предусловия завершены с ошибкой.')
        for skipped in data.skipped_test_cases:
            doc.add_paragraph(f'  {skipped.test_case_id}, Шаг {skipped.step_number}. {skipped.description}')

    doc = _add_decision_solution(doc, data)

    if data.failed_tests == 0 and not data.has_blocking_defect:
        p = doc.add_paragraph()
        p.add_run(f'Все тест-кейсы, включенные в прогон {data.run_name} версии\n')
        p.add_run(f'{data.run_version} завершены успешно. Релиз может быть рекомендован для установки\n')
        p.add_run('в части проверенного функционала.')

    elif data.has_blocking_defect:
        p = doc.add_paragraph()
        p.add_run('Требуется доработка релиза (устранение блокирующего дефекта) и\n')
        p.add_run('проведение повторного автоматизированного тестирования.')

    elif data.critical_errors:
        p = doc.add_paragraph()
        p.add_run('Требуется доработка релиза в части критичного функционала:\n')
        for err in data.critical_errors:
            doc.add_paragraph(f'🗹 - {err.test_case_id}, Шаг {err.step_number}. {err.description}')
        for err in data.non_critical_errors[:3]:
            doc.add_paragraph(f'◻ - {err.test_case_id}, Шаг {err.step_number}. {err.description}')
        doc.add_paragraph('\nи проведение повторного автоматизированного тестирования.')

    elif data.non_critical_errors:
        p = doc.add_paragraph()
        p.add_run('Рекомендуется доработка релиза в части **не**критичного функционала:\n')
        for err in data.non_critical_errors:
            doc.add_paragraph(f'- {err.test_case_id}, Шаг {err.step_number}. {err.description}')
        doc.add_paragraph('\nРелиз может быть рекомендован для установки в части успешно проверенного функционала.')

    doc = _add_author_of_tests_section(doc, data)

    doc.add_page_break()
    doc.add_heading('Приложение', 1)

    doc.add_paragraph(f'Прогон: {data.run_name}')
    doc.add_paragraph(f'Дата и время: {data.run_datetime.strftime("%H:%M:%S %d.%m.%Y")}')

    if data.has_blocking_defect:
        status = "блокирующий дефект"
    elif data.failed_tests > 0:
        status = "с ошибкой"
    else:
        status = "Успешно"
    doc.add_paragraph(f'Статус: {status}')
    doc.add_paragraph(
        f'Инициатор запуска прогона: {data.initiator_name if data.initiator_name else "____________________"}')

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['№ п/п', 'Название Тест-кейса', 'Статус\nУспешно/с ошибкой\n/пропущен', 'Длительность выполнения']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
            apply_times_new_roman(cell.paragraphs[0])

    for idx, tc in enumerate(data.test_cases_table, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = tc.name
        row.cells[2].text = tc.status.value
        row.cells[3].text = tc.duration

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                apply_times_new_roman(paragraph)

    filename = f"Сводный_отчет_{data.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)

    return filepath

def generate_short_report(data: TestRunData) -> str:
    os.makedirs(REPORTS_DIR, exist_ok=True)

    doc = Document()
    apply_times_new_roman_to_doc(doc)

    title = doc.add_heading('Краткий отчёт', 0)
    apply_times_new_roman(title)

    doc = _add_metrics_section(doc, data)

    doc = _add_failed_tests_section(doc, data)

    doc = _add_skipped_tests_section(doc, data)

    doc = _add_decision_solution(doc, data)

    if data.failed_tests == 0 and not data.has_blocking_defect:
        p = doc.add_paragraph()
        p.add_run(f'Все тест-кейсы, включенные в прогон {data.run_name} версии\n')
        p.add_run(f'{data.run_version} завершены успешно. Релиз может быть рекомендован для установки\n')
        p.add_run('в части проверенного функционала.')

    elif data.has_blocking_defect:
        p = doc.add_paragraph()
        p.add_run('Требуется доработка релиза (устранение блокирующего дефекта) и\n')
        p.add_run('проведение повторного автоматизированного тестирования.')

    elif data.critical_errors:
        p = doc.add_paragraph()
        p.add_run('Требуется доработка релиза в части критичного функционала\n')
        p.add_run('и проведение повторного автоматизированного тестирования.')

    elif data.non_critical_errors:
        p = doc.add_paragraph()
        p.add_run('Рекомендуется устранить **не**критичные ошибки. При этом Релиз может\n')
        p.add_run('быть рекомендован для установки в части успешно проверенного\n')
        p.add_run('функционала.')

    doc = _add_author_of_tests_section(doc, data)

    filename = f"Краткий_отчет_{data.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)

    return filepath


def generate_full_report(data: TestRunData) -> str:
    os.makedirs(REPORTS_DIR, exist_ok=True)

    doc = Document()
    apply_times_new_roman_to_doc(doc)

    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ОТЧЁТ О ТЕСТИРОВАНИИ')
    run.bold = True
    run.font.size = Pt(16)
    apply_times_new_roman(p)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.project_name)
    run.bold = True
    run.font.size = Pt(14)
    apply_times_new_roman(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'версия {data.version}')
    run.bold = True
    apply_times_new_roman(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026')
    run.bold = True
    apply_times_new_roman(p)

    doc.add_page_break()

    doc.add_heading('СОДЕРЖАНИЕ', level=1)

    doc.add_paragraph('1 ВВЕДЕНИЕ')
    doc.add_paragraph('2 ПРЕДМЕТ ТЕСТИРОВАНИЯ')
    doc.add_paragraph('3 ОБЛАСТЬ ОХВАТЫВАЕМЫХ ВОПРОСОВ')
    doc.add_paragraph('4 ОГРАНИЧЕНИЯ ТЕСТИРОВАНИЯ')
    doc.add_paragraph('5 ТЕСТИРОВАНИЕ НОВОЙ ФУНКЦИОНАЛЬНОСТИ')
    doc.add_paragraph('5.1.1 Тестирование изменений')
    doc.add_paragraph('5.1.2 Тестирование нового функционала')
    doc.add_paragraph('6 ИНТЕГРАЦИОННОЕ ТЕСТИРОВАНИЕ')
    doc.add_paragraph('7 АВТОМАТИЗИРОВАННОЕ РЕГРЕССИОННОЕ ТЕСТИРОВАНИЕ')
    doc.add_paragraph('8 ИТОГИ ТЕСТИРОВАНИЯ')
    doc.add_paragraph('9 ВЫВОДЫ И РЕКОМЕНДАЦИИ')

    doc.add_page_break()

    doc.add_heading('1 ВВЕДЕНИЕ', level=1)
    doc.add_paragraph('Период проведения тестирования:')
    doc.add_paragraph(
        f'{data.run_datetime.strftime("%d.%m.%Y %H:%M")} - {data.run_datetime.strftime("%d.%m.%Y %H:%M")}.')
    doc.add_paragraph()
    doc.add_paragraph('Место проведения тестирования:')
    doc.add_paragraph(
        f'Тестирование {data.project_name} проводится на ресурсах заказчика с использованием ресурсов и программного обеспечения.')
    doc.add_paragraph()
    doc.add_paragraph('Цели проведения тестирования:')
    doc.add_paragraph(data.run_name)

    doc.add_page_break()

    doc.add_heading('2 ПРЕДМЕТ ТЕСТИРОВАНИЯ', level=1)
    doc.add_paragraph(f'Предметом тестирования является {data.project_name} версия {data.version}.')
    doc.add_paragraph(data.project_name)

    doc.add_page_break()

    doc.add_heading('3 ОБЛАСТЬ ОХВАТЫВАЕМЫХ ВОПРОСОВ', level=1)
    doc.add_paragraph('1. Тестирование новой функциональности:')
    doc.add_paragraph(f'    1.  Общее количество изменений, {len(data.failed_test_cases)}')
    doc.add_paragraph(f'    2.  Количество нового функционала, {len(data.test_cases_table)}.')
    doc.add_paragraph('Подробная информация о проверках приведена в разделе «Тестирование новой функциональности».')
    doc.add_paragraph()
    doc.add_paragraph('2. Интеграционное тестирование:')
    doc.add_paragraph(
        f'Перечень программного обеспечения, с которым проводилось интеграционное тестирование {data.project_name}.')
    doc.add_paragraph('Подробная информация о проверках приведена в разделе «Интеграционное тестирование».')
    doc.add_paragraph()
    doc.add_paragraph('3. Автоматизированное регрессионное тестирование:')
    doc.add_paragraph(f'Автоматизированное регрессионное тестирование проводилось в объеме {data.total_tests}.')
    doc.add_paragraph(
        'Подробная информация о проверках приведена в разделе «Автоматизированное регрессионное тестирование».')

    doc.add_page_break()

    doc.add_heading('4 ОГРАНИЧЕНИЯ ТЕСТИРОВАНИЯ', level=1)
    doc.add_paragraph(f'Тестирование {data.project_name} проводилось с рядом ограничений:')
    doc.add_paragraph(
        f'- проводилось ограниченное тестирование функционала в объеме, указанном в прогоне {data.run_name} версии {data.run_version}, другой функционал не тестировался;')
    doc.add_paragraph('- проводилось тестирование через пользовательский интерфейс методом «черного ящика».')

    doc.add_page_break()

    doc.add_heading('5 ТЕСТИРОВАНИЕ НОВОЙ ФУНКЦИОНАЛЬНОСТИ', level=1)

    doc.add_heading('Тестирование изменений', level=2)
    doc.add_paragraph('Результаты тестирования изменений:')

    table = doc.add_table(rows=len(data.test_cases_table) + 2, cols=4)
    table.style = 'Table Grid'
    headers = ['№ п/п', 'Название тест-кейса', 'Время запуска', 'Статус']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    table.cell(1,
               1).text = 'В данной таблице присутствуют тест-кейсы, прошедшие рефакторинг (определяется по версии тест-кейса) и включенных на замену подобным тест-кейсам из предыдущего прогона'
    table.cell(1, 3).text = 'Успешно / С ошибками'

    for idx, tc in enumerate(data.test_cases_table, 1):
        row = idx + 1
        table.cell(row, 0).text = str(idx)
        table.cell(row, 1).text = tc.name
        table.cell(row, 2).text = data.run_datetime.strftime("%H:%M:%S")
        table.cell(row, 3).text = tc.status.value

    doc.add_paragraph(f'Всего выполнено проверок: {len(data.test_cases_table)}')
    doc.add_paragraph(f'Выполнено с ошибками: {data.failed_tests}')

    doc.add_heading('Тестирование нового функционала', level=2)
    doc.add_paragraph('Результаты тестирования нового функционала:')

    table2 = doc.add_table(rows=len(data.test_cases_table) + 2, cols=4)
    table2.style = 'Table Grid'
    for i, header in enumerate(headers):
        table2.cell(0, i).text = header

    table2.cell(1,
                1).text = 'В данной таблице присутствуют тест-кейсы, добавленные в текущий прогон и отсутствующие в предыдущем прогоне'
    table2.cell(1, 3).text = 'Успешно / С ошибками'

    for idx, tc in enumerate(data.test_cases_table, 1):
        row = idx + 1
        table2.cell(row, 0).text = str(idx)
        table2.cell(row, 1).text = tc.name
        table2.cell(row, 2).text = data.run_datetime.strftime("%H:%M:%S")
        table2.cell(row, 3).text = tc.status.value

    doc.add_paragraph(f'Выполнено проверок: {len(data.test_cases_table)}.')
    doc.add_paragraph(f'Выполнено с ошибками: {data.failed_tests}')

    doc.add_page_break()

    doc.add_heading('6 ИНТЕГРАЦИОННОЕ ТЕСТИРОВАНИЕ', level=1)
    doc.add_paragraph('Результаты интеграционного тестирования:')

    table3 = doc.add_table(rows=len(data.test_cases_table) + 2, cols=4)
    table3.style = 'Table Grid'
    for i, header in enumerate(headers):
        table3.cell(0, i).text = header

    table3.cell(1,
                1).text = 'В данной таблице присутствуют тест-кейсы из других проектов (такое может быть, если тест-кейс в предусловии имеет выполнение ТК из другого проекта'
    table3.cell(1, 3).text = 'Успешно / С ошибками'

    for idx, tc in enumerate(data.test_cases_table, 1):
        row = idx + 1
        table3.cell(row, 0).text = str(idx)
        table3.cell(row, 1).text = tc.name
        table3.cell(row, 2).text = data.run_datetime.strftime("%H:%M:%S")
        table3.cell(row, 3).text = tc.status.value

    doc.add_paragraph(f'Выполнено проверок: {len(data.test_cases_table)}.')
    doc.add_paragraph(f'Выполнено с ошибками: {data.failed_tests}')

    doc.add_page_break()

    doc.add_heading('7 АВТОМАТИЗИРОВАННОЕ РЕГРЕССИОННОЕ ТЕСТИРОВАНИЕ', level=1)
    doc.add_paragraph('Результаты автоматизированного регрессионного тестирования:')

    table4 = doc.add_table(rows=len(data.test_cases_table) + 2, cols=4)
    table4.style = 'Table Grid'
    for i, header in enumerate(headers):
        table4.cell(0, i).text = header

    table4.cell(1, 1).text = 'В данной таблице присутствуют тест-кейсы идентичные предыдущему прогону'
    table4.cell(1, 3).text = 'Успешно / С ошибками'

    for idx, tc in enumerate(data.test_cases_table, 1):
        row = idx + 1
        table4.cell(row, 0).text = str(idx)
        table4.cell(row, 1).text = tc.name
        table4.cell(row, 2).text = data.run_datetime.strftime("%H:%M:%S")
        table4.cell(row, 3).text = tc.status.value

    doc.add_paragraph(f'Выполнено проверок: {data.total_tests}.')
    doc.add_paragraph(f'Выполнено с ошибками: {data.failed_tests}')

    doc.add_page_break()

    doc.add_heading('8 ИТОГИ ТЕСТИРОВАНИЯ', level=1)

    if data.failed_tests == 0 and not data.has_blocking_defect:
        doc.add_paragraph()
        doc.add_paragraph('По результатам тестирования ошибок не обнаружено.')

    elif data.has_blocking_defect and data.blocking_defect:
        doc.add_paragraph()
        doc.add_paragraph(
            f'В период тестирования выявлен блокирующий дефект при прохождении тест-кейса {data.blocking_defect.test_case_id} на шаге №{data.blocking_defect.step_number}')
        doc.add_paragraph()

        table5 = doc.add_table(rows=2, cols=4)
        table5.style = 'Table Grid'
        headers5 = ['Шаг', 'Описание действия', 'Ожидаемый результат', 'Фактический результат']
        for i, header in enumerate(headers5):
            table5.cell(0, i).text = header
            for run in table5.cell(0, i).paragraphs[0].runs:
                run.bold = True
        table5.cell(1, 0).text = str(data.blocking_defect.step_number)
        table5.cell(1, 1).text = data.blocking_defect.description
        table5.cell(1, 2).text = 'Ожидаемый результат'
        table5.cell(1, 3).text = 'Фактический результат'

    elif data.failed_test_cases:
        doc.add_paragraph()
        doc.add_paragraph(f'По результатам тестирования выявлены {len(data.failed_test_cases)} ошибок:')
        doc.add_paragraph()
        doc.add_paragraph('  ------------------------------------------------------------------------')
        doc.add_paragraph('        ID тест-кейса   Название тест-кейса     Шаг    Описание действия')
        doc.add_paragraph('  ----- --------------- ---------------------- -------- ------------------')

        for idx, err in enumerate(data.failed_test_cases[:5]):
            mark = "🗹" if err in data.critical_errors else "◻"
            doc.add_paragraph(
                f'    {mark}        {err.test_case_id:<15} {err.description[:30]:<22} {err.step_number:<8} {err.description}')

        doc.add_paragraph('  ------------------------------------------------------------------------')
        doc.add_paragraph()
        doc.add_paragraph(
            f'Из них {len(data.critical_errors)} с критичным функционалом и {len(data.non_critical_errors)} с некритичным функционалом.')

    doc.add_page_break()

    doc.add_heading('9 ВЫВОДЫ И РЕКОМЕНДАЦИИ', level=1)

    if data.failed_tests == 0 and not data.has_blocking_defect:
        doc.add_paragraph()
        doc.add_paragraph('Выводы:')
        doc.add_paragraph(f'1.  В ходе тестирования {data.project_name} версии {data.version} ошибки не выявлены.')
        doc.add_paragraph(
            f'2.  В части протестированных функций {data.project_name} версии {data.version} работоспособно.')
        doc.add_paragraph()
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph('1.  Рекомендации отсутствуют.')

    elif data.has_blocking_defect:
        doc.add_paragraph()
        doc.add_paragraph('Выводы:')
        doc.add_paragraph(
            f'1.  В ходе тестирования {data.project_name} версии {data.version} выявлен блокирующий дефект.')
        doc.add_paragraph(
            f'2.  Провести тестирование всех функций {data.project_name} версии {data.version} не представляется возможным.')
        doc.add_paragraph()
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph('1.  Необходимо доработать релиз (устранить блокирующий дефект).')
        doc.add_paragraph('2.  Провести повторное автоматизированное тестирование.')

    elif data.critical_errors:
        doc.add_paragraph()
        doc.add_paragraph('Выводы:')
        doc.add_paragraph(
            f'1.  В ходе тестирования {data.project_name} версии {data.version} выявлено {len(data.failed_test_cases)} ошибок, из них:')
        doc.add_paragraph(f'    - критичных -- {len(data.critical_errors)};')
        doc.add_paragraph(f'    - некритичных - {len(data.non_critical_errors)}.')
        doc.add_paragraph()
        doc.add_paragraph(
            f'2.  В части протестированных функций {data.project_name} версии {data.version} требует доработки релиза в части критичного функционала.')
        doc.add_paragraph()
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph(
            '1.  Необходимо доработать релиз (устранить критические ошибки, указанные в результатах настоящего Отчета).')
        doc.add_paragraph('2.  Провести повторное автоматизированное тестирование.')

    elif data.non_critical_errors and not data.critical_errors:
        doc.add_paragraph()
        doc.add_paragraph('Выводы:')
        doc.add_paragraph(
            f'1.  В ходе тестирования {data.project_name} версии {data.version} выявлено {len(data.failed_test_cases)} ошибок, из них:')
        doc.add_paragraph(f'    - критичных -- 0;')
        doc.add_paragraph(f'    - некритичных - {len(data.non_critical_errors)}.')
        doc.add_paragraph()
        doc.add_paragraph(
            f'2.  В части протестированных функций {data.project_name} версии {data.version} в целом работоспособно.')
        doc.add_paragraph()
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph('1.  Релиз может быть установлен в части успешно проверенного функционала.')
        doc.add_paragraph(
            '2.  В плановом порядке устранить некритичные ошибки, указанные в результатах настоящего Отчета.')

    doc = _add_author_of_tests_section(doc,data)

    filename = f"Полный_отчет_{data.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)

    return filepath